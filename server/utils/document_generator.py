"""
Document generation utilities for Word and PDF formats.
Supports briefings, letters, memos, and meeting notes.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from typing import List
from datetime import datetime


def create_briefing_doc(title: str, summary: str, key_points: List[str], 
                       action_items: List[str], output_path: str, format: str = "docx"):
    """
    Create a briefing document with structured sections.
    """
    if format == "docx":
        _create_briefing_docx(title, summary, key_points, action_items, output_path)
    elif format == "pdf":
        _create_briefing_pdf(title, summary, key_points, action_items, output_path)
    else:
        raise ValueError(f"Unsupported format: {format}")


def _create_briefing_docx(title: str, summary: str, key_points: List[str], 
                          action_items: List[str], output_path: str):
    """Create briefing as Word document."""
    doc = Document()
    
    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # Spacer
    
    # Summary section
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(summary)
    
    doc.add_paragraph()  # Spacer
    
    # Key Points section
    doc.add_heading("Key Points", level=1)
    for i, point in enumerate(key_points, 1):
        doc.add_paragraph(f"{i}. {point}")
    
    doc.add_paragraph()  # Spacer
    
    # Action Items section
    doc.add_heading("Action Items", level=1)
    for i, item in enumerate(action_items, 1):
        para = doc.add_paragraph(style='List Bullet')
        para.text = item
    
    doc.save(output_path)


def _create_briefing_pdf(title: str, summary: str, key_points: List[str], 
                        action_items: List[str], output_path: str):
    """Create briefing as PDF document."""
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=12,
        spaceBefore=12
    )
    
    # Title
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Date
    date_text = f"Date: {datetime.now().strftime('%B %d, %Y')}"
    story.append(Paragraph(date_text, styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    
    # Summary
    story.append(Paragraph("Executive Summary", heading_style))
    story.append(Paragraph(summary, styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Key Points
    story.append(Paragraph("Key Points", heading_style))
    for i, point in enumerate(key_points, 1):
        story.append(Paragraph(f"{i}. {point}", styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Action Items
    story.append(Paragraph("Action Items", heading_style))
    for item in action_items:
        story.append(Paragraph(f"• {item}", styles['Normal']))
    
    doc.build(story)


def create_document(doc_type: str, title: str, content: str, 
                   output_path: str, format: str = "docx"):
    """
    Create a formatted document (letter, memo, meeting_notes).
    """
    if format == "docx":
        _create_document_docx(doc_type, title, content, output_path)
    elif format == "pdf":
        _create_document_pdf(doc_type, title, content, output_path)
    else:
        raise ValueError(f"Unsupported format: {format}")


def _create_document_docx(doc_type: str, title: str, content: str, output_path: str):
    """Create document as Word format."""
    doc = Document()
    
    if doc_type == "letter":
        # Letter format
        doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
        doc.add_paragraph()
        doc.add_paragraph(title)
        doc.add_paragraph()
        
        # Content paragraphs
        for para in content.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.add_paragraph()
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("[Signature]")
        
    elif doc_type == "memo":
        # Memo format
        doc.add_heading("MEMORANDUM", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        doc.add_paragraph(f"TO: [Recipient]")
        doc.add_paragraph(f"FROM: [Sender]")
        doc.add_paragraph(f"DATE: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"RE: {title}")
        doc.add_paragraph()
        doc.add_paragraph("_" * 60)
        doc.add_paragraph()
        
        # Content
        for para in content.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
    elif doc_type == "meeting_notes":
        # Meeting notes format
        doc.add_heading(title, level=0)
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph()
        
        doc.add_heading("Attendees:", level=2)
        doc.add_paragraph("[List attendees]")
        doc.add_paragraph()
        
        doc.add_heading("Notes:", level=2)
        for para in content.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.add_paragraph()
        doc.add_heading("Action Items:", level=2)
        doc.add_paragraph("[List action items]")
        
    else:
        # Generic document
        doc.add_heading(title, level=0)
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph()
        
        for para in content.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
    
    doc.save(output_path)


def _create_document_pdf(doc_type: str, title: str, content: str, output_path: str):
    """Create document as PDF format."""
    doc_pdf = SimpleDocTemplate(output_path, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=20,
        alignment=TA_CENTER if doc_type == "memo" else TA_LEFT
    )
    
    if doc_type == "letter":
        # Letter format
        story.append(Paragraph(datetime.now().strftime('%B %d, %Y'), styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(title, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        for para in content.split('\n\n'):
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("Sincerely,", styles['Normal']))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph("[Signature]", styles['Normal']))
        
    elif doc_type == "memo":
        # Memo format
        story.append(Paragraph("MEMORANDUM", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph("TO: [Recipient]", styles['Normal']))
        story.append(Paragraph("FROM: [Sender]", styles['Normal']))
        story.append(Paragraph(f"DATE: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Paragraph(f"RE: {title}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("_" * 80, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        for para in content.split('\n\n'):
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
    elif doc_type == "meeting_notes":
        # Meeting notes format
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph("Attendees:", styles['Heading2']))
        story.append(Paragraph("[List attendees]", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph("Notes:", styles['Heading2']))
        for para in content.split('\n\n'):
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("Action Items:", styles['Heading2']))
        story.append(Paragraph("[List action items]", styles['Normal']))
        
    else:
        # Generic document
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        for para in content.split('\n\n'):
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
    
    doc_pdf.build(story)
